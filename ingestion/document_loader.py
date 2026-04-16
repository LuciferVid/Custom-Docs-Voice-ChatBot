import os
import pdfplumber
import PyPDF2
from docx import Document
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def load_pdf(file_path: str) -> list[dict]:
    pages = []
    doc_name = os.path.basename(file_path)
    
    # Strategy 1: pdfplumber (High precision)
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    pages.append({
                        "text": page_text.strip(),
                        "page_number": i + 1,
                        "source_file": doc_name
                    })
    except Exception as e:
        logger.error(f"pdfplumber failed for {file_path}: {e}")

    # Strategy 2: PyPDF2 (Failover)
    if not pages:
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages.append({
                            "text": page_text.strip(),
                            "page_number": i + 1,
                            "source_file": doc_name
                        })
        except Exception as e:
            logger.error(f"PyPDF2 failed for {file_path}: {e}")

    return pages

def load_document(file_path: str) -> list[dict]:
    """
    Loads text from PDF, DOCX, TXT, or MD with failover strategy.
    Returns list of dicts (pages).
    """
    ext = os.path.splitext(file_path)[1].lower()
    doc_name = os.path.basename(file_path)
    pages = []
    
    try:
        if ext == ".pdf":
            pages = load_pdf(file_path)
        elif ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            if full_text.strip():
                pages.append({"text": full_text, "page_number": 1, "source_file": doc_name})
        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
                if text.strip():
                    pages.append({"text": text, "page_number": 1, "source_file": doc_name})
        
        return pages
    except Exception as e:
        logger.error(f"Load error for {file_path}: {e}")
        return []

def load_docx(file_path: str) -> list[dict]:
    """
    Extracts text from a Word document (.docx).
    Groups paragraphs into chunks of ~500 characters.
    """
    pages = []
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        content = "\n".join(full_text)
        # Split into chunks of ~500 chars
        chunk_size = 500
        for i in range(0, len(content), chunk_size):
            chunk = content[i:i + chunk_size]
            if len(chunk.strip()) >= 50:
                pages.append({
                    "text": chunk.strip(),
                    "page_number": "N/A",
                    "source_file": os.path.basename(file_path)
                })
    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
    return pages

def load_txt(file_path: str) -> list[dict]:
    """
    Reads text content from a .txt or .md file.
    Splits into chunks of 500 characters.
    """
    pages = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        chunk_size = 500
        for i in range(0, len(content), chunk_size):
            chunk = content[i:i + chunk_size]
            if len(chunk.strip()) >= 50:
                pages.append({
                    "text": chunk.strip(),
                    "page_number": "N/A",
                    "source_file": os.path.basename(file_path)
                })
    except Exception as e:
        logger.error(f"Error loading TXT {file_path}: {e}")
    return pages

def load_document(file_path: str) -> list[dict]:
    """
    Detects file type and routes to the correct loader.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext in [".txt", ".md"]:
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def get_supported_extensions() -> list[str]:
    """
    Returns list of supported extensions.
    """
    return [".pdf", ".docx", ".txt", ".md"]
